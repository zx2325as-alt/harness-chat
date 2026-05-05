from __future__ import annotations

import csv
import io
import os
from dataclasses import dataclass, asdict
from typing import Any, Dict, List, Optional

MAX_TEXT_CHARS = 180_000
CHUNK_SIZE = 6_000
CHUNK_OVERLAP = 800
MAX_PDF_PAGES = 80
MAX_SHEET_ROWS = 2_000

_DOC_LIMITS: Dict[str, int] = {
    "max_text_chars": MAX_TEXT_CHARS,
    "chunk_size": CHUNK_SIZE,
    "chunk_overlap": CHUNK_OVERLAP,
    "max_pdf_pages": MAX_PDF_PAGES,
    "max_sheet_rows": MAX_SHEET_ROWS,
}


def configure_document_limits(cfg: Optional[Dict[str, Any]]) -> None:
    """由 app 启动时从 harness.documents 注入；缺省沿用模块常量。"""
    if not cfg or not isinstance(cfg, dict):
        return
    m = dict(_DOC_LIMITS)
    if cfg.get("max_text_chars") is not None:
        try:
            m["max_text_chars"] = max(10_000, int(cfg["max_text_chars"]))
        except (TypeError, ValueError):
            pass
    if cfg.get("chunk_size") is not None:
        try:
            m["chunk_size"] = max(512, int(cfg["chunk_size"]))
        except (TypeError, ValueError):
            pass
    if cfg.get("chunk_overlap") is not None:
        try:
            m["chunk_overlap"] = max(0, int(cfg["chunk_overlap"]))
        except (TypeError, ValueError):
            pass
    if cfg.get("max_pdf_pages") is not None:
        try:
            m["max_pdf_pages"] = max(1, int(cfg["max_pdf_pages"]))
        except (TypeError, ValueError):
            pass
    if cfg.get("max_sheet_rows") is not None:
        try:
            m["max_sheet_rows"] = max(64, int(cfg["max_sheet_rows"]))
        except (TypeError, ValueError):
            pass
    _DOC_LIMITS.clear()
    _DOC_LIMITS.update(m)

TEXT_DOCUMENT_EXTS = {
    "txt",
    "md",
    "markdown",
    "json",
    "yaml",
    "yml",
    "toml",
    "ini",
    "cfg",
    "conf",
    "env",
    "log",
    "xml",
    "html",
    "htm",
    "css",
    "scss",
    "less",
    "js",
    "mjs",
    "cjs",
    "jsx",
    "ts",
    "tsx",
    "py",
    "java",
    "kt",
    "kts",
    "scala",
    "go",
    "rs",
    "c",
    "cc",
    "cpp",
    "cxx",
    "h",
    "hh",
    "hpp",
    "cs",
    "php",
    "rb",
    "swift",
    "sh",
    "bash",
    "zsh",
    "ps1",
    "sql",
    "r",
    "lua",
    "pl",
    "pm",
    "proto",
    "properties",
    "gradle",
    "vue",
    "svelte",
    "dart",
}
SPECIAL_TEXT_FILENAMES = {
    "dockerfile",
    "makefile",
    "jenkinsfile",
    ".env",
    ".gitignore",
    ".npmrc",
    ".yarnrc",
    ".editorconfig",
    ".prettierrc",
    ".eslintrc",
    "cmakelists.txt",
}
SUPPORTED_DOCUMENT_EXTS = TEXT_DOCUMENT_EXTS | {
    "csv",
    "pdf",
    "doc",
    "docx",
    "xls",
    "xlsx",
} | SPECIAL_TEXT_FILENAMES


@dataclass
class ExtractedDocument:
    name: str
    ext: str
    status: str
    content: str
    chunks: List[Dict[str, Any]]
    error: str | None = None
    meta: Dict[str, Any] | None = None

    def to_dict(self) -> Dict[str, Any]:
        data = asdict(self)
        data["meta"] = data["meta"] or {}
        return data


def _chunk_text(
    text: str, chunk_size: Optional[int] = None, overlap: Optional[int] = None
) -> List[Dict[str, Any]]:
    chunk_size = int(chunk_size if chunk_size is not None else _DOC_LIMITS["chunk_size"])
    overlap = int(overlap if overlap is not None else _DOC_LIMITS["chunk_overlap"])
    text = (text or "").strip()
    if not text:
        return []
    ov = max(0, min(overlap, chunk_size // 2))
    step = max(1, chunk_size - ov)
    chunks = []
    for idx, start in enumerate(range(0, len(text), step), start=1):
        part = text[start : start + chunk_size]
        if not part:
            break
        chunks.append({"index": idx, "content": part, "start": start, "end": start + len(part)})
    return chunks


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def detect_document_ext(filename: str) -> str:
    name = os.path.basename(filename or "").strip()
    lower_name = name.lower()
    if lower_name in SPECIAL_TEXT_FILENAMES:
        return lower_name
    return os.path.splitext(lower_name)[1].lower().lstrip(".")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("缺少 pypdf 依赖，无法解析 PDF。") from exc

    reader = PdfReader(io.BytesIO(data))
    pages = []
    limit_pages = _DOC_LIMITS["max_pdf_pages"]
    for page_no, page in enumerate(reader.pages, start=1):
        if page_no > limit_pages:
            pages.append(f"【已截断：仅解析前 {limit_pages} 页】")
            break
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"【第 {page_no} 页】\n{text}")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except Exception as exc:
        raise RuntimeError("缺少 python-docx 依赖，无法解析 DOCX。") from exc

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table_idx, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            parts.append(f"【表格 {table_idx}】\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as exc:
        raise RuntimeError("缺少 openpyxl 依赖，无法解析 XLSX。") from exc

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        lines = [f"【工作表：{ws.title}】"]
        cap = _DOC_LIMITS["max_sheet_rows"]
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if row_idx > cap:
                lines.append(f"【已截断：仅解析前 {cap} 行】")
                break
            values = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in values):
                lines.append(" | ".join(values))
        parts.append("\n".join(lines))
    return "\n\n".join(parts)


def _extract_xls(data: bytes) -> str:
    try:
        import xlrd
    except Exception as exc:
        raise RuntimeError("缺少 xlrd 依赖，无法解析 XLS。") from exc

    book = xlrd.open_workbook(file_contents=data)
    parts = []
    for sheet in book.sheets():
        lines = [f"【工作表：{sheet.name}】"]
        cap = _DOC_LIMITS["max_sheet_rows"]
        for row_idx in range(min(sheet.nrows, cap)):
            values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
            if any(v.strip() for v in values):
                lines.append(" | ".join(values))
        if sheet.nrows > cap:
            lines.append(f"【已截断：仅解析前 {cap} 行】")
        parts.append("\n".join(lines))
    return "\n\n".join(parts)


def _extract_csv(data: bytes) -> str:
    text = _decode_text(data)
    reader = csv.reader(io.StringIO(text))
    lines = [" | ".join(row) for row in reader if any(cell.strip() for cell in row)]
    return "\n".join(lines)


def extract_document(filename: str, data: bytes) -> ExtractedDocument:
    name = filename or "未命名文件"
    ext = detect_document_ext(name)
    try:
        if ext in TEXT_DOCUMENT_EXTS or ext in SPECIAL_TEXT_FILENAMES or ext == "csv":
            content = _extract_csv(data) if ext == "csv" else _decode_text(data)
        elif ext == "pdf":
            content = _extract_pdf(data)
        elif ext == "docx":
            content = _extract_docx(data)
        elif ext == "doc":
            content = _decode_text(data)
        elif ext == "xlsx":
            content = _extract_xlsx(data)
        elif ext == "xls":
            content = _extract_xls(data)
        else:
            raise RuntimeError(f"暂不支持的文件格式：.{ext or 'unknown'}")

        content = (content or "").strip()
        truncated = False
        max_chars = _DOC_LIMITS["max_text_chars"]
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True
        if not content:
            raise RuntimeError("文件中没有提取到可用文本。")

        chunks = _chunk_text(content)
        return ExtractedDocument(
            name=name,
            ext=ext,
            status="ok",
            content=content,
            chunks=chunks,
            meta={"chars": len(content), "chunks": len(chunks), "truncated": truncated},
        )
    except Exception as exc:
        return ExtractedDocument(
            name=name,
            ext=ext,
            status="error",
            content="",
            chunks=[],
            error=str(exc),
            meta={},
        )
